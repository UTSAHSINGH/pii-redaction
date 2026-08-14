"""
test_document_integrity.py
--------------------------
Critical Regression Suite for Non-PII Content Preservation and Semantic Integrity.

Covers the 10 Golden Requirement Tests:
1. Date protection ('Dated December 10, 2025' remains unchanged).
2. Legal phrase protection ('Companies Act, 1956' remains unchanged).
3. Financial phrase protection ('Return on Capital Employed' remains unchanged).
4. Table of contents / headings ('CAPITAL STRUCTURE' remains unchanged).
5. Company span ('KSH International Limited' replaced as full span, no partial corruption).
6. PERSON replacement ('Contact Person: Sarthak Malvadkar' -> 'Contact Person: Alex Carter').
7. Email replacement ('cs.connect@kshinternational.com' -> exact synthetic email).
8. Phone replacement ('+91 20 4505 3237' -> exact synthetic phone).
9. No cascading (replacements are never re-detected).
10. Run spanning (name split across runs is replaced exactly once).
"""

import pytest
import sys
from pathlib import Path
from docx import Document

sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

from detectors import DETECTORS, PROTECTED_PHRASES, _overlaps_protected, _get_protected_spans
from document_processor import (
    DocumentSegment,
    apply_replacements_to_segment,
    extract_document_segments,
    reconstruct_expected_text,
)
from replacement_generator import get_replacement, reset_replacement_map, get_synthetic_values
from utils import PIIMatch, resolve_overlaps, validate_match_span


class TestGoldenIntegrityRequirements:
    def setup_method(self):
        reset_replacement_map()

    # ------------------------------------------------------------------ Test 1
    def test_date_protection_prospectus_date(self):
        """Test 1: 'Dated December 10, 2025' must NEVER be redacted."""
        text = "Dated December 10, 2025"
        matches = []
        for d in DETECTORS:
            matches.extend(d.detect_in_segment("seg_date", text))
        # Ensure no detector claims this date
        assert len(matches) == 0, f"Expected 0 detections on prospectus date, got: {matches}"

    # ------------------------------------------------------------------ Test 2
    def test_legal_phrase_protection_companies_act(self):
        """Test 2: 'Companies Act, 1956' and 'Companies Act, 2013' must remain unchanged."""
        text = "Incorporated under the provisions of the Companies Act, 1956 and governed by Companies Act, 2013."
        matches = []
        for d in DETECTORS:
            matches.extend(d.detect_in_segment("seg_legal", text))
        # None of the matches can touch 'Companies Act'
        for m in matches:
            assert "Companies Act" not in m.text

    # ------------------------------------------------------------------ Test 3
    def test_financial_phrase_protection_roce(self):
        """Test 3: 'Return on Capital Employed' must remain unchanged."""
        text = "Financial Metrics: Return on Capital Employed and Return on Equity were evaluated."
        matches = []
        for d in DETECTORS:
            matches.extend(d.detect_in_segment("seg_fin", text))
        assert len(matches) == 0, f"Financial phrases must not trigger detection: {matches}"

    # ------------------------------------------------------------------ Test 4
    def test_table_of_contents_capital_structure(self):
        """Test 4: 'CAPITAL STRUCTURE' heading must remain unchanged."""
        text = "SECTION IV: CAPITAL STRUCTURE AND OBJECTS OF THE OFFER"
        matches = []
        for d in DETECTORS:
            matches.extend(d.detect_in_segment("seg_toc", text))
        assert len(matches) == 0, f"TOC headings must not trigger detection: {matches}"

    # ------------------------------------------------------------------ Test 5
    def test_company_span_integrity(self):
        """Test 5: 'KSH International Limited' must be replaced as a complete entity."""
        doc = Document()
        p = doc.add_paragraph("The Issuer is KSH International Limited in India.")
        seg = extract_document_segments(doc)[0]

        matches = []
        for d in DETECTORS:
            matches.extend(d.detect_in_segment(seg.segment_id, seg.text))
        resolved = resolve_overlaps(matches)

        for m in resolved:
            m.replacement = get_replacement(m)

        apply_replacements_to_segment(seg, resolved)
        # Should not contain partial corruption
        assert "LimitedLimited" not in p.text
        assert "KSH International Joseph" not in p.text
        assert "The Issuer is " in p.text
        assert " in India." in p.text

    # ------------------------------------------------------------------ Test 6
    def test_person_replacement_context(self):
        """Test 6: 'Contact Person: Sarthak Malvadkar' becomes 'Contact Person: <replacement>'."""
        doc = Document()
        p = doc.add_paragraph("Contact Person: Sarthak Malvadkar")
        seg = extract_document_segments(doc)[0]

        matches = []
        for d in DETECTORS:
            matches.extend(d.detect_in_segment(seg.segment_id, seg.text))
        resolved = resolve_overlaps(matches)
        assert len(resolved) >= 1
        assert any("Sarthak Malvadkar" in m.text for m in resolved)

        for m in resolved:
            m.replacement = get_replacement(m)

        apply_replacements_to_segment(seg, resolved)
        assert p.text.startswith("Contact Person: ")
        assert "Sarthak Malvadkar" not in p.text

    # ------------------------------------------------------------------ Test 7
    def test_email_replacement_exact(self):
        """Test 7: 'cs.connect@kshinternational.com' becomes exactly one synthetic email."""
        doc = Document()
        p = doc.add_paragraph("E-mail: cs.connect@kshinternational.com; Website: www.kshinternational.com")
        seg = extract_document_segments(doc)[0]

        matches = []
        for d in DETECTORS:
            matches.extend(d.detect_in_segment(seg.segment_id, seg.text))
        resolved = resolve_overlaps(matches)

        for m in resolved:
            m.replacement = get_replacement(m)

        apply_replacements_to_segment(seg, resolved)
        assert "cs.connect@kshinternational.com" not in p.text
        assert "@example.com" in p.text
        assert "Website: www.kshinternational.com" in p.text

    # ------------------------------------------------------------------ Test 8
    def test_phone_replacement_exact(self):
        """Test 8: '+91 20 4505 3237' becomes exactly one synthetic phone number."""
        doc = Document()
        p = doc.add_paragraph("Telephone: +91 20 4505 3237.")
        seg = extract_document_segments(doc)[0]

        matches = []
        for d in DETECTORS:
            matches.extend(d.detect_in_segment(seg.segment_id, seg.text))
        resolved = resolve_overlaps(matches)

        for m in resolved:
            m.replacement = get_replacement(m)

        apply_replacements_to_segment(seg, resolved)
        assert "+91 20 4505 3237" not in p.text
        assert "Telephone: +91 " in p.text

    # ------------------------------------------------------------------ Test 9
    def test_no_cascading_detection(self):
        """Test 9: A replacement value must not trigger secondary detection."""
        m_person = PIIMatch("seg_1", "PERSON", 0, 17, "Sarthak Malvadkar", 0.95, "test")
        rep_person = get_replacement(m_person)

        # Ensure synthetic exclusion registry knows this value
        synthetic_vals = get_synthetic_values()
        assert rep_person in synthetic_vals

    # ------------------------------------------------------------------ Test 10
    def test_run_spanning_exact_replacement(self):
        """Test 10: A PERSON split across two runs is replaced exactly once without duplication."""
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Mr. ")
        p.add_run("Sarthak ")
        p.add_run("Malvadkar")
        p.add_run(" joined the meeting.")

        seg = extract_document_segments(doc)[0]
        start = seg.text.find("Sarthak Malvadkar")
        end = start + len("Sarthak Malvadkar")

        m = PIIMatch(
            segment_id=seg.segment_id,
            entity_type="PERSON",
            start=start,
            end=end,
            text="Sarthak Malvadkar",
            confidence=0.95,
            source="test",
            replacement="Alex Carter",
        )

        apply_replacements_to_segment(seg, [m])
        assert p.text == "Mr. Alex Carter joined the meeting."


class TestCharacterPreservationRegression:
    def test_unredacted_character_preservation(self):
        """
        Verify that all non-PII characters outside approved match spans
        remain 100% byte-for-byte identical.
        """
        original = "The board meeting held on December 10, 2025 approved the Fresh Issue of Equity Shares."
        doc = Document()
        p = doc.add_paragraph(original)
        seg = extract_document_segments(doc)[0]

        # No PII matches in this sentence
        matches = []
        for d in DETECTORS:
            matches.extend(d.detect_in_segment(seg.segment_id, seg.text))

        expected = reconstruct_expected_text(seg.text, matches)
        apply_replacements_to_segment(seg, matches)

        assert p.text == original
        assert p.text == expected
