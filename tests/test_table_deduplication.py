"""
test_table_deduplication.py
---------------------------
Tests asserting that table cells and merged rows are deduplicated by underlying XML element identity.
Prevents duplicate replacement stacking on merged cells.
"""

import sys
from pathlib import Path
import pytest
from docx import Document

sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

from document_processor import extract_document_segments


def test_table_cell_deduplication():
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    # Merge cells in first row
    cell_a = table.cell(0, 0)
    cell_b = table.cell(0, 1)
    cell_a.merge(cell_b)
    cell_a.paragraphs[0].text = "Merged Header: Sarthak Malvadkar"

    # Row 2 cells
    table.cell(1, 0).paragraphs[0].text = "Cell 1,0"
    table.cell(1, 1).paragraphs[0].text = "Cell 1,1"

    segments = extract_document_segments(doc)
    # Total segments should be 3 (1 merged cell paragraph + 2 row 2 paragraphs), NOT 4
    merged_segs = [s for s in segments if "Merged Header" in s.text]
    assert len(merged_segs) == 1, "Merged table cell paragraphs must be deduplicated to exactly 1 segment!"
