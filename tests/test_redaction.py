"""
test_redaction.py
-----------------
Tests for replacement generation, deterministic mapping, collision-freedom,
case preservation, and right-to-left run-level DOCX replacement.
"""

import pytest
import sys
from pathlib import Path
from docx import Document

sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

from replacement_generator import (
    get_replacement,
    get_replacement_map,
    reset_replacement_map,
)
from utils import PIIMatch, resolve_overlaps, luhn_checksum
from document_processor import (
    DocumentSegment,
    apply_replacements_to_segment,
    extract_document_segments,
    reconstruct_expected_text,
)


class TestReplacementGenerator:
    def setup_method(self):
        reset_replacement_map()

    def test_deterministic_consistency(self):
        m1 = PIIMatch("seg_1", "PERSON", 0, 17, "Sarthak Malvadkar", 0.95, "test")
        m2 = PIIMatch("seg_2", "PERSON", 5, 22, "Sarthak Malvadkar", 0.95, "test")
        rep1 = get_replacement(m1)
        rep2 = get_replacement(m2)
        assert rep1 == rep2, "Same original must always produce identical replacement"

    def test_collision_free(self):
        m1 = PIIMatch("seg_1", "PERSON", 0, 17, "Sarthak Malvadkar", 0.95, "test")
        m2 = PIIMatch("seg_2", "PERSON", 0, 21, "Kushal Subbayya Hegde", 0.95, "test")
        rep1 = get_replacement(m1)
        rep2 = get_replacement(m2)
        assert rep1 != rep2, "Different original entities must receive different replacements"

    def test_email_format(self):
        m = PIIMatch("seg_1", "EMAIL", 0, 31, "cs.connect@kshinternational.com", 0.99, "test")
        rep = get_replacement(m)
        assert "@example.com" in rep

    def test_phone_format(self):
        m = PIIMatch("seg_1", "PHONE", 0, 16, "+91 20 4505 3237", 0.95, "test")
        rep = get_replacement(m)
        assert rep.startswith("+91 ")

    def test_all_caps_preservation_without_suffix_corruption(self):
        # ALL-CAPS input should become ALL-CAPS output without trailing letters
        m = PIIMatch("seg_1", "PERSON", 0, 21, "KUSHAL SUBBAYYA HEGDE", 0.95, "test")
        rep = get_replacement(m)
        assert rep.isupper(), f"Expected ALL-CAPS, got: {rep}"
        assert not rep.endswith("ter"), "Must not have trailing mixed-case artifact"

    def test_credit_card_valid_luhn(self):
        m = PIIMatch("seg_1", "CREDIT_CARD", 0, 19, "4111-1111-1111-1111", 0.99, "test")
        rep = get_replacement(m)
        assert luhn_checksum(rep.replace("-", "").replace(" ", ""))


class TestRunLevelReplacement:
    def test_single_run_replacement(self):
        doc = Document()
        p = doc.add_paragraph("Contact Person: Sarthak Malvadkar for queries.")
        seg = extract_document_segments(doc)[0]

        match = PIIMatch(
            segment_id=seg.segment_id,
            entity_type="PERSON",
            start=16,
            end=33,
            text="Sarthak Malvadkar",
            confidence=0.95,
            source="test",
            replacement="Alex Carter",
        )

        apply_replacements_to_segment(seg, [match])
        assert p.text == "Contact Person: Alex Carter for queries."

    def test_multi_run_spanning_replacement(self):
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Contact: ")
        p.add_run("Sarthak ")
        p.add_run("Malvadkar")
        p.add_run(" (CS)")

        seg = extract_document_segments(doc)[0]
        # Text is "Contact: Sarthak Malvadkar (CS)"
        # "Sarthak Malvadkar" spans run 1 and run 2
        start = seg.text.find("Sarthak Malvadkar")
        end = start + len("Sarthak Malvadkar")

        match = PIIMatch(
            segment_id=seg.segment_id,
            entity_type="PERSON",
            start=start,
            end=end,
            text="Sarthak Malvadkar",
            confidence=0.95,
            source="test",
            replacement="Alex Carter",
        )

        apply_replacements_to_segment(seg, [match])
        assert p.text == "Contact: Alex Carter (CS)"

    def test_multiple_replacements_in_same_segment_right_to_left(self):
        doc = Document()
        p = doc.add_paragraph("Email cs.connect@kshinternational.com or call +91 20 4505 3237 now.")
        seg = extract_document_segments(doc)[0]

        m_email = PIIMatch(
            segment_id=seg.segment_id,
            entity_type="EMAIL",
            start=6,
            end=37,
            text="cs.connect@kshinternational.com",
            confidence=0.99,
            source="test",
            replacement="alex.carter@example.com",
        )
        m_phone = PIIMatch(
            segment_id=seg.segment_id,
            entity_type="PHONE",
            start=46,
            end=62,
            text="+91 20 4505 3237",
            confidence=0.95,
            source="test",
            replacement="+91 91234 56789",
        )

        apply_replacements_to_segment(seg, [m_email, m_phone])
        expected = "Email alex.carter@example.com or call +91 91234 56789 now."
        assert p.text == expected


class TestOverlapResolution:
    def test_structured_pii_takes_priority_over_ner(self):
        m_email = PIIMatch("seg_1", "EMAIL", 10, 35, "john.smith@nuvama.com", 0.99, "regex")
        m_person = PIIMatch("seg_1", "PERSON", 10, 20, "john.smith", 0.85, "ner")

        resolved = resolve_overlaps([m_person, m_email])
        assert len(resolved) == 1
        assert resolved[0].entity_type == "EMAIL"

    def test_non_overlapping_both_kept(self):
        m1 = PIIMatch("seg_1", "PERSON", 0, 10, "Alex Smith", 0.95, "test")
        m2 = PIIMatch("seg_1", "EMAIL", 15, 35, "alex@example.com", 0.99, "test")
        resolved = resolve_overlaps([m1, m2])
        assert len(resolved) == 2
