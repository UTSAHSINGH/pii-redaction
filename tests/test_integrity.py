"""
test_integrity.py
-----------------
Tests for cryptographic-grade document integrity verification.
Asserts that every character outside approved PII spans is byte-for-byte preserved.
"""

import sys
from pathlib import Path
import pytest
from docx import Document

sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

from integrity_verifier import apply_only_approved_spans, verify_document_integrity
from document_processor import apply_replacements_to_segment, extract_document_segments
from utils import PIIMatch, save_json


def test_apply_only_approved_spans_exact():
    original = "Contact Person: Sarthak Malvadkar, Company Secretary; Telephone: +91 20 4505 3237"
    approved = [
        {"start": 16, "end": 33, "text": "Sarthak Malvadkar", "replacement": "Logan Sullivan"},
        {"start": 65, "end": 81, "text": "+91 20 4505 3237", "replacement": "+91 91234 56789"},
    ]
    expected = "Contact Person: Logan Sullivan, Company Secretary; Telephone: +91 91234 56789"
    result = apply_only_approved_spans(original, approved)
    assert result == expected


def test_independent_verifier_pass(tmp_path):
    orig_doc = Document()
    p = orig_doc.add_paragraph("Dated December 10, 2025. Contact Sarthak Malvadkar for queries.")
    orig_path = tmp_path / "orig.docx"
    orig_doc.save(str(orig_path))

    red_doc = Document(str(orig_path))
    segs = extract_document_segments(red_doc)
    match = PIIMatch(
        segment_id=segs[0].segment_id,
        entity_type="PERSON",
        start=33,
        end=50,
        text="Sarthak Malvadkar",
        confidence=0.95,
        source="test",
        replacement="Logan Sullivan",
    )
    apply_replacements_to_segment(segs[0], [match])
    red_path = tmp_path / "red.docx"
    red_doc.save(str(red_path))

    log_path = tmp_path / "log.json"
    save_json([{
        "segment_id": segs[0].segment_id,
        "entity_type": "PERSON",
        "text": "Sarthak Malvadkar",
        "start": 33,
        "end": 50,
        "replacement": "Logan Sullivan",
    }], log_path)

    report = verify_document_integrity(orig_path, red_path, log_path)
    assert report["status"] == "PASS"
    assert report["segments_with_unexpected_changes"] == 0
    assert report["unauthorized_changed_characters"] == 0
    assert report["unauthorized_change_rate"] == 0.0
    assert report["non_pii_preservation_rate"] == 1.0


def test_independent_verifier_detects_unauthorized_mutation(tmp_path):
    orig_doc = Document()
    orig_doc.add_paragraph("The Offer is 100% Book Built. Contact Sarthak Malvadkar.")
    orig_path = tmp_path / "orig.docx"
    orig_doc.save(str(orig_path))

    red_doc = Document(str(orig_path))
    # Unauthorized corruption: 'The Offer' -> 'Quinn Hayes'
    red_doc.paragraphs[0].text = "Quinn Hayes is 100% Book Built. Contact Logan Sullivan."
    red_path = tmp_path / "red.docx"
    red_doc.save(str(red_path))

    log_path = tmp_path / "log.json"
    # Approved log only covers Sarthak Malvadkar
    save_json([{
        "segment_id": "para_0",
        "entity_type": "PERSON",
        "text": "Sarthak Malvadkar",
        "start": 38,
        "end": 55,
        "replacement": "Logan Sullivan",
    }], log_path)

    report = verify_document_integrity(orig_path, red_path, log_path)
    assert report["status"] == "FAIL"
    assert report["segments_with_unexpected_changes"] > 0
    assert report["unauthorized_changed_characters"] > 0
    assert report["unauthorized_change_rate"] > 0.0
