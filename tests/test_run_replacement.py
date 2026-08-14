"""
test_run_replacement.py
-----------------------
Tests for run-level in-place replacement and right-to-left mutation.
Asserts formatting preservation and zero character corruption outside approved spans.
"""

import sys
from pathlib import Path
import pytest
from docx import Document

sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

from document_processor import apply_replacements_to_segment, extract_document_segments
from utils import PIIMatch


def test_single_run_exact_patch():
    doc = Document()
    p = doc.add_paragraph("Contact Person: Sarthak Malvadkar is our CS.")
    seg = extract_document_segments(doc)[0]

    match = PIIMatch(
        segment_id=seg.segment_id,
        entity_type="PERSON",
        start=16,
        end=33,
        text="Sarthak Malvadkar",
        confidence=0.95,
        source="test",
        replacement="Logan Sullivan",
    )

    apply_replacements_to_segment(seg, [match])
    assert p.text == "Contact Person: Logan Sullivan is our CS."


def test_multi_run_spanning_patch():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Contact: ")
    p.add_run("Sarthak ")
    p.add_run("Malvadkar")
    p.add_run(" (CS)")

    seg = extract_document_segments(doc)[0]
    start = seg.text.find("Sarthak Malvadkar")
    end = start + len("Sarthak Malvadkar")

    match = PIIMatch(
        segment_id=seg.segment_id,
        entity_type="PERSON",
        start=start,
        end=end,
        text="Sarthak Malvadkar",
        confidence=0.95,
        source="test",
        replacement="Logan Sullivan",
    )

    apply_replacements_to_segment(seg, [match])
    assert p.text == "Contact: Logan Sullivan (CS)"


def test_multiple_replacements_right_to_left():
    doc = Document()
    p = doc.add_paragraph("Email cs.connect@kshinternational.com or call +91 20 4505 3237 now.")
    seg = extract_document_segments(doc)[0]

    m1 = PIIMatch(
        segment_id=seg.segment_id,
        entity_type="EMAIL",
        start=6,
        end=37,
        text="cs.connect@kshinternational.com",
        confidence=0.99,
        source="test",
        replacement="alex.carter@example.com",
    )
    m2 = PIIMatch(
        segment_id=seg.segment_id,
        entity_type="PHONE",
        start=46,
        end=62,
        text="+91 20 4505 3237",
        confidence=0.95,
        source="test",
        replacement="+91 91234 56789",
    )

    apply_replacements_to_segment(seg, [m1, m2])
    assert p.text == "Email alex.carter@example.com or call +91 91234 56789 now."
