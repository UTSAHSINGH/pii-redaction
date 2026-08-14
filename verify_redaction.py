import sys
from docx import Document
sys.path.insert(0, 'src')

doc = Document('output/redacted_output.docx')
print('Redacted DOCX stats:')
print(f'  Paragraphs: {len(doc.paragraphs)}')
print(f'  Tables: {len(doc.tables)}')
print()

original_emails = [
    'cs.connect@kshinternational.com', 
    'ksh.ipo@nuvama.com', 
    'siddharth.jadhav@hdfcbank.com', 
    'ipo@trilegal.com'
]
original_names = ['Sarthak Malvadkar', 'Kushal Subbayya Hegde']

full_text = ' '.join(p.text for p in doc.paragraphs)
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            full_text += ' ' + cell.text

print('Checking original PII is NOT in redacted output:')
for email in original_emails:
    present = email.lower() in full_text.lower()
    status = 'STILL PRESENT - FAIL' if present else 'REDACTED OK - PASS'
    print(f'  {email}: {status}')

print()
for name in original_names:
    present = name.lower() in full_text.lower()
    status = 'STILL PRESENT - FAIL' if present else 'REDACTED OK - PASS'
    print(f'  {name}: {status}')

print()
print('Replacement emails visible in output:')
import re
fake_emails = re.findall(r'[a-z.]+@example\.com', full_text)[:5]
for e in fake_emails:
    print(f'  {e}')

print()
print('DOCX opens successfully: PASS')
