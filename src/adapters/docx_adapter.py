"""
docx_adapter.py
---------------
High-precision Microsoft Word (.docx) document adapter.
Preserves styles, run boundaries, hyperlinks, tables, and XML part deduplication.
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import Any, Dict, List, Set, Tuple
import uuid

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from adapters.base import DocumentAdapter, NormalizedDocument
from models import DocumentSegment, PIIMatch


class DocxAdapter(DocumentAdapter):
    """Adapter for .docx documents."""

    supported_extensions = [".docx", ".docm"]

    def extract(self, file_source: str | Path | bytes, filename: str = "document.docx") -> NormalizedDocument:
        if isinstance(file_source, bytes):
            stream = io.BytesIO(file_source)
            raw_bytes = file_source
        else:
            p = Path(file_source)
            raw_bytes = p.read_bytes()
            stream = io.BytesIO(raw_bytes)

        doc = Document(stream)
        segments: List[DocumentSegment] = []
        seen_element_ids: Set[int] = set()

        # 1. Extract Body Paragraphs
        for idx, p in enumerate(doc.paragraphs):
            elem_id = id(p._element)
            if elem_id in seen_element_ids:
                continue
            seen_element_ids.add(elem_id)
            seg_id = f"para_{idx}"
            segments.append(DocumentSegment(
                segment_id=seg_id,
                text=p.text,
                metadata={"type": "paragraph", "index": idx},
                element_ref=p,
            ))

        # 2. Extract Table Cells (deduplicating merged cells)
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    for p_idx, p in enumerate(cell.paragraphs):
                        elem_id = id(p._element)
                        if elem_id in seen_element_ids:
                            continue
                        seen_element_ids.add(elem_id)
                        seg_id = f"table_{t_idx}_r{r_idx}_c{c_idx}_p{p_idx}"
                        segments.append(DocumentSegment(
                            segment_id=seg_id,
                            text=p.text,
                            metadata={"type": "table_cell", "table": t_idx, "row": r_idx, "col": c_idx},
                            element_ref=p,
                        ))

        # 3. Extract Headers and Footers (deduplicating shared XML parts)
        processed_parts: Set[str] = set()
        for s_idx, section in enumerate(doc.sections):
            for hf_type, hf in [
                ("header", section.header),
                ("footer", section.footer),
                ("first_page_header", section.first_page_header),
                ("first_page_footer", section.first_page_footer),
                ("even_page_header", section.even_page_header),
                ("even_page_footer", section.even_page_footer),
            ]:
                if not hf or hf.is_linked_to_previous:
                    continue
                part_name = getattr(hf.part, "partname", str(id(hf)))
                if part_name in processed_parts:
                    continue
                processed_parts.add(part_name)

                for p_idx, p in enumerate(hf.paragraphs):
                    elem_id = id(p._element)
                    if elem_id in seen_element_ids:
                        continue
                    seen_element_ids.add(elem_id)
                    seg_id = f"section_{s_idx}_{hf_type}_p{p_idx}"
                    segments.append(DocumentSegment(
                        segment_id=seg_id,
                        text=p.text,
                        metadata={"type": hf_type, "section": s_idx},
                        element_ref=p,
                    ))

        return NormalizedDocument(
            document_id=str(uuid.uuid4()),
            filename=filename,
            file_type="docx",
            segments=segments,
            raw_bytes=raw_bytes,
            metadata={
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "sections": len(doc.sections),
            },
            native_doc=doc,
        )

    def apply_redactions(
        self,
        doc: NormalizedDocument,
        approved_matches: List[PIIMatch],
    ) -> bytes:
        native_doc: Document = doc.native_doc
        matches_by_seg: Dict[str, List[PIIMatch]] = {}
        for m in approved_matches:
            matches_by_seg.setdefault(m.segment_id, []).append(m)

        for seg in doc.segments:
            seg_matches = matches_by_seg.get(seg.segment_id, [])
            if not seg_matches:
                continue

            para = seg.element_ref
            if not para or not hasattr(para, "runs"):
                continue

            # Sort matches descending by start position (right-to-left)
            sorted_matches = sorted(seg_matches, key=lambda m: m.start, reverse=True)

            # Build character-to-run map from existing runs
            char_map = self._build_char_to_run_map(para.runs)

            for match in sorted_matches:
                start = match.start
                end = match.end
                replacement = match.replacement or ""

                if start >= len(char_map) or end > len(char_map):
                    continue

                start_run_idx, start_offset = char_map[start]
                end_run_idx, end_offset = char_map[end - 1]

                if start_run_idx == end_run_idx:
                    # Single run replacement
                    run = para.runs[start_run_idx]
                    old_text = run.text
                    run.text = old_text[:start_offset] + replacement + old_text[end_offset + 1:]
                else:
                    # Multi-run spanning replacement
                    first_run = para.runs[start_run_idx]
                    last_run = para.runs[end_run_idx]

                    first_run.text = first_run.text[:start_offset] + replacement
                    for mid_idx in range(start_run_idx + 1, end_run_idx):
                        para.runs[mid_idx].text = ""
                    last_run.text = last_run.text[end_offset + 1:]

                # Rebuild char map after mutation
                char_map = self._build_char_to_run_map(para.runs)

        # Patch hyperlinks
        self._patch_hyperlinks(native_doc, approved_matches)

        out_stream = io.BytesIO()
        native_doc.save(out_stream)
        return out_stream.getvalue()

    def _build_char_to_run_map(self, runs) -> List[Tuple[int, int]]:
        char_map = []
        for run_idx, run in enumerate(runs):
            for offset in range(len(run.text)):
                char_map.append((run_idx, offset))
        return char_map

    def _patch_hyperlinks(self, doc: Document, matches: List[PIIMatch]) -> None:
        email_replacements = {m.text.strip(): m.replacement for m in matches if m.entity_type == "EMAIL" and m.replacement}
        if not email_replacements:
            return
        for rel in doc.part.rels.values():
            if rel.reltype == RT.HYPERLINK:
                target = rel.target_ref
                if target.startswith("mailto:"):
                    raw_email = target[len("mailto:"):].split("?")[0].strip()
                    if raw_email in email_replacements:
                        rel._target = f"mailto:{email_replacements[raw_email]}"
