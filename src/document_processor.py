"""
document_processor.py
---------------------
Immutable Document Model and Run-Level DOCX Processing Engine.

Key Principles:
1. Every paragraph, table cell paragraph, and unique header/footer is extracted as a
   structured DocumentSegment with an immutable original text representation.
2. Detection runs strictly against segment.text.
3. Replacements are applied in-place to the underlying Run objects within each segment
   using a strictly right-to-left order, preventing offset drift and format loss.
4. Hyperlink relationship targets (mailto:) are updated consistently.
5. Unique XML headers and footers are deduplicated so no part is processed multiple times.
"""

from __future__ import annotations

import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple
from xml.etree import ElementTree as ET

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from utils import PIIMatch, setup_logger

log = setup_logger("document_processor")


# ---------------------------------------------------------------------------
# PseudoRun for Hyperlink Text Elements
# ---------------------------------------------------------------------------

class _PseudoRun:
    """
    Wraps text elements within a w:hyperlink so they expose the same interface
    as standard python-docx Run objects (r.text getter and setter).
    """

    def __init__(self, r_elem, t_elem) -> None:
        self._r = r_elem
        self._t = t_elem

    @property
    def text(self) -> str:
        return self._t.text or ""

    @text.setter
    def text(self, value: str) -> None:
        self._t.text = value


# ---------------------------------------------------------------------------
# Document Segment Model
# ---------------------------------------------------------------------------

@dataclass
class DocumentSegment:
    """
    Structured, self-contained unit of document text with run references.
    """

    segment_id: str                   # Unique identifier (e.g. "para_12", "table_2_r_1_c_0_p_0")
    container_type: str               # "paragraph" | "table_cell" | "header" | "footer"
    location: str                     # Human-readable location description
    text: str                         # Immutable original text of this segment
    runs: List[object]                # Ordered list of Run and _PseudoRun objects
    paragraph_ref: Paragraph          # Reference to underlying python-docx Paragraph


def _extract_runs_from_paragraph(paragraph: Paragraph) -> List[object]:
    """
    Extract all text-bearing Run and _PseudoRun objects in visual document order,
    including runs nested inside <w:hyperlink> elements.
    """
    runs: List[object] = []
    p_elem = paragraph._element

    for child in p_elem:
        tag = child.tag
        if tag.endswith("r"):
            # Standard run
            runs.append(Run(child, paragraph))
        elif tag.endswith("hyperlink"):
            # Hyperlink container: extract nested w:r -> w:t
            for r_child in child:
                if r_child.tag.endswith("r"):
                    for t_child in r_child:
                        if t_child.tag.endswith("t"):
                            runs.append(_PseudoRun(r_child, t_child))

    return runs


# ---------------------------------------------------------------------------
# Document Segment Extraction
# ---------------------------------------------------------------------------

def extract_document_segments(doc: Document) -> List[DocumentSegment]:
    """
    Extract all text segments from the document in reading order:
    1. Body paragraphs
    2. Table cell paragraphs (each cell paragraph processed exactly once)
    3. Unique section headers
    4. Unique section footers
    """
    segments: List[DocumentSegment] = []

    # 1. Body Paragraphs
    for p_idx, para in enumerate(doc.paragraphs):
        runs = _extract_runs_from_paragraph(para)
        text = "".join(r.text or "" for r in runs)
        seg = DocumentSegment(
            segment_id=f"para_{p_idx}",
            container_type="paragraph",
            location=f"Body Paragraph {p_idx}",
            text=text,
            runs=runs,
            paragraph_ref=para,
        )
        segments.append(seg)

    # 2. Table Cells (iterate rows -> cells -> paragraphs, deduplicated by underlying XML element ID)
    seen_cell_para_elems: Set[int] = set()
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, para in enumerate(cell.paragraphs):
                    elem_id = id(para._element)
                    if elem_id not in seen_cell_para_elems:
                        seen_cell_para_elems.add(elem_id)
                        runs = _extract_runs_from_paragraph(para)
                        text = "".join(r.text or "" for r in runs)
                        seg = DocumentSegment(
                            segment_id=f"table_{t_idx}_r_{r_idx}_c_{c_idx}_p_{p_idx}",
                            container_type="table_cell",
                            location=f"Table {t_idx} Row {r_idx} Cell {c_idx} Para {p_idx}",
                            text=text,
                            runs=runs,
                            paragraph_ref=para,
                        )
                        segments.append(seg)

    # 3. Headers (deduplicated by underlying XML element ID)
    seen_header_elems: Set[int] = set()
    for s_idx, section in enumerate(doc.sections):
        for hdr_type, hdr in [
            ("header", section.header),
            ("first_page_header", section.first_page_header),
            ("even_page_header", section.even_page_header),
        ]:
            if hdr is not None and hasattr(hdr, "_element") and hdr._element is not None:
                elem_id = id(hdr._element)
                if elem_id not in seen_header_elems:
                    seen_header_elems.add(elem_id)
                    try:
                        for p_idx, para in enumerate(hdr.paragraphs):
                            runs = _extract_runs_from_paragraph(para)
                            text = "".join(r.text or "" for r in runs)
                            seg = DocumentSegment(
                                segment_id=f"sec_{s_idx}_{hdr_type}_p_{p_idx}",
                                container_type="header",
                                location=f"Section {s_idx} {hdr_type} Para {p_idx}",
                                text=text,
                                runs=runs,
                                paragraph_ref=para,
                            )
                            segments.append(seg)
                    except Exception as exc:
                        log.debug("Header extraction exception: %s", exc)

    # 4. Footers (deduplicated by underlying XML element ID)
    seen_footer_elems: Set[int] = set()
    for s_idx, section in enumerate(doc.sections):
        for ftr_type, ftr in [
            ("footer", section.footer),
            ("first_page_footer", section.first_page_footer),
            ("even_page_footer", section.even_page_footer),
        ]:
            if ftr is not None and hasattr(ftr, "_element") and ftr._element is not None:
                elem_id = id(ftr._element)
                if elem_id not in seen_footer_elems:
                    seen_footer_elems.add(elem_id)
                    try:
                        for p_idx, para in enumerate(ftr.paragraphs):
                            runs = _extract_runs_from_paragraph(para)
                            text = "".join(r.text or "" for r in runs)
                            seg = DocumentSegment(
                                segment_id=f"sec_{s_idx}_{ftr_type}_p_{p_idx}",
                                container_type="footer",
                                location=f"Section {s_idx} {ftr_type} Para {p_idx}",
                                text=text,
                                runs=runs,
                                paragraph_ref=para,
                            )
                            segments.append(seg)
                    except Exception as exc:
                        log.debug("Footer extraction exception: %s", exc)

    log.info("Extracted %d document segments across paragraphs, tables, and headers/footers.", len(segments))
    return segments


# ---------------------------------------------------------------------------
# Run-Level Replacement Engine (Right-to-Left)
# ---------------------------------------------------------------------------

def apply_replacements_to_segment(
    segment: DocumentSegment,
    matches: List[PIIMatch],
) -> None:
    """
    Apply accepted PII replacements to the Run objects in this segment.
    """
    if not matches or not segment.runs:
        return

    # Sort matches from right to left
    sorted_matches = sorted(matches, key=lambda m: m.start, reverse=True)

    for match in sorted_matches:
        if match.replacement is None:
            continue

        m_start = match.start
        m_end = match.end
        rep = match.replacement

        # Compute current run character boundaries
        run_spans: List[Tuple[int, int, object]] = []
        cur_pos = 0
        for run in segment.runs:
            rlen = len(run.text or "")
            run_spans.append((cur_pos, cur_pos + rlen, run))
            cur_pos += rlen

        # Find overlapping runs for [m_start, m_end)
        overlapping = [
            (r_start, r_end, run)
            for r_start, r_end, run in run_spans
            if r_start < m_end and r_end > m_start
        ]

        if not overlapping:
            continue

        if len(overlapping) == 1:
            # Single run replacement
            r_start, r_end, run = overlapping[0]
            local_start = max(0, m_start - r_start)
            local_end = min(len(run.text or ""), m_end - r_start)
            old_text = run.text or ""
            run.text = old_text[:local_start] + rep + old_text[local_end:]

        else:
            # Multi-run replacement
            # First overlapping run gets prefix + full replacement
            first_r_start, first_r_end, first_run = overlapping[0]
            local_start = max(0, m_start - first_r_start)
            prefix = (first_run.text or "")[:local_start]
            first_run.text = prefix + rep

            # Middle overlapping runs are cleared
            for _, _, mid_run in overlapping[1:-1]:
                mid_run.text = ""

            # Last overlapping run keeps suffix after match
            last_r_start, last_r_end, last_run = overlapping[-1]
            local_end = min(len(last_run.text or ""), m_end - last_r_start)
            suffix = (last_run.text or "")[local_end:]
            last_run.text = suffix


# ---------------------------------------------------------------------------
# Expected Text Reconstruction (For Regression / Integrity Verification)
# ---------------------------------------------------------------------------

def reconstruct_expected_text(original_text: str, matches: List[PIIMatch]) -> str:
    """
    Reconstruct the expected text of a segment by applying replacements right-to-left
    on the immutable original text string.
    """
    if not matches:
        return original_text

    sorted_matches = sorted(matches, key=lambda m: m.start, reverse=True)
    text = original_text
    for m in sorted_matches:
        if m.replacement is not None:
            text = text[:m.start] + m.replacement + text[m.end:]
    return text


# ---------------------------------------------------------------------------
# Hyperlink Relationship Target Patching
# ---------------------------------------------------------------------------

def patch_hyperlink_targets(doc: Document, matches: List[PIIMatch]) -> int:
    """
    Find any hyperlink relationships (e.g. mailto: links) containing redacted email
    addresses and update their Target URLs to match the synthetic replacement.
    """
    email_map = {
        m.text.strip().lower(): m.replacement
        for m in matches
        if m.entity_type == "EMAIL" and m.replacement
    }
    if not email_map:
        return 0

    patched_count = 0
    # Search document part relationships
    for part in [doc.part] + [s.header.part for s in doc.sections if s.header] + [s.footer.part for s in doc.sections if s.footer]:
        try:
            for rel_id, rel in part.rels.items():
                if rel.is_external and rel.target_ref:
                    target = rel.target_ref
                    for orig_email, rep_email in email_map.items():
                        if orig_email in target.lower():
                            new_target = target.lower().replace(orig_email, rep_email)
                            rel.target_ref = new_target
                            patched_count += 1
                            log.info("Patched hyperlink rel %s: '%s' → '%s'", rel_id, target, new_target)
        except Exception as exc:
            log.debug("Hyperlink patching check exception: %s", exc)

    return patched_count


# ---------------------------------------------------------------------------
# Document Statistics
# ---------------------------------------------------------------------------

def get_document_stats(doc: Document) -> dict:
    """Return counts of structural document elements."""
    unique_headers: Set[int] = set()
    unique_footers: Set[int] = set()

    for s in doc.sections:
        for hdr in [s.header, s.first_page_header, s.even_page_header]:
            if hdr is not None and hasattr(hdr, "_element") and hdr._element is not None:
                unique_headers.add(id(hdr._element))
        for ftr in [s.footer, s.first_page_footer, s.even_page_footer]:
            if ftr is not None and hasattr(ftr, "_element") and ftr._element is not None:
                unique_footers.add(id(ftr._element))

    return {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "sections": len(doc.sections),
        "unique_headers": len(unique_headers),
        "unique_footers": len(unique_footers),
    }
